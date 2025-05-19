import os
import re
import time
from PyPDF2 import PdfReader
from docx import Document
import google.generativeai as genai
import streamlit as st
import pandas as pd
import io
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import requests # Make sure this is imported if used in URL form
from bs4 import BeautifulSoup # Make sure this is imported if used in URL form

# --- Streamlit Page Configuration ---
st.set_page_config(layout="wide")

st.markdown(
    """
    <style>
    .title {
        text-align: center;
        font-size: 2.5em;
        color: #007bff; /* Bleu Bootstrap */
    }
    .subtitle {
        text-align: center;
        font-size: 1.5em;
        color: #6c757d; /* Gris Bootstrap */
    }
    .stTextInput > label,
    .stTextArea > label,
    .stFileUploader > div > div:first-child {
        text-align: left !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown("<h1 class='title'>Application de Sélection de CV</h1>", unsafe_allow_html=True)
st.markdown("<h2 class='subtitle'>Optimisez votre processus de recrutement avec l'IA</h2>", unsafe_allow_html=True)

# --- Initialize Session State ---
if 'job_title' not in st.session_state:
    st.session_state['job_title'] = ""
if 'job_experience' not in st.session_state:
    st.session_state['job_experience'] = ""
if 'job_description' not in st.session_state:
    st.session_state['job_description'] = ""
if 'all_extracted_resume_texts' not in st.session_state:
    st.session_state.all_extracted_resume_texts = []
if 'knowledge_base_content' not in st.session_state:
    st.session_state.knowledge_base_content = ""
if 'chat_messages' not in st.session_state:
    st.session_state.chat_messages = []
if 'gemini_chat_model_instance' not in st.session_state:
    st.session_state.gemini_chat_model_instance = None
if 'gemini_chat_session_object' not in st.session_state:
    st.session_state.gemini_chat_session_object = None
if 'api_key_configured' not in st.session_state:
    st.session_state.api_key_configured = False

# --- API Key Configuration Function ---
def attempt_api_configuration(keys_to_try):
    """
    Attempts to configure the genai API with a list of keys.
    Returns the model instance on success, None on failure.
    Sets st.session_state.api_key_configured accordingly.
    """
    idx = 0
    st.session_state.api_key_configured = False # Reset before trying
    while idx < len(keys_to_try):
        try:
            key = keys_to_try[idx]
            genai.configure(api_key=key)
            # Test the configuration by trying to create a model instance
            model = genai.GenerativeModel('gemini-1.5-flash')
            st.session_state.api_key_configured = True
            st.info(f"API configurée avec succès (utilisation de la clé index {idx}).")
            return model # Return the model instance
        except Exception as e:
            st.warning(f"Échec de configuration de l'API avec la clé index {idx}: {e}")
            idx += 1
            if "quota" in str(e).lower() or "permission" in str(e).lower():
                st.warning(f"Problème potentiel de quota ou de permission avec la clé index {idx-1}.")
            # Add more specific error handling if needed (e.g., invalid API key format)
    st.error("Toutes les clés API ont échoué ou un autre problème de configuration est survenu. Le service est temporairement indisponible.")
    st.session_state.api_key_configured = False
    return None

# --- UI for Job Details Input ---
input_method = st.radio("Méthode de saisie des détails du poste:", ["Lien URL (Sauf IFCARJOB Lien)", "Saisir manuellement"], horizontal=True)

if input_method == "Lien URL (Sauf IFCARJOB Lien)":
    with st.form(key="url_form"):
        url = st.text_input("Entrez l'URL pour extraire les détails du poste:")
        submit_url = st.form_submit_button("Extraire les détails du poste")
        if submit_url and url:
            try:
                headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
                response = requests.get(url, headers=headers, timeout=10)
                response.raise_for_status()
                soup = BeautifulSoup(response.content, "html.parser")
                
                # Attempt to find title - more robustly
                title_tag = soup.find('h1')
                if not title_tag: title_tag = soup.find('title') # Fallback to <title> tag
                st.session_state['job_title'] = title_tag.text.strip() if title_tag else "Titre du poste non trouvé (vérifiez URL/sélecteurs)"
                
                # Example selectors - THESE WILL LIKELY NEED ADJUSTMENT PER WEBSITE
                # For experience, you might need to search for keywords like "expérience", "experience", "profil"
                experience_element = soup.find(lambda tag: tag.name == 'div' and ('card-wrapper-inner' in tag.get('class', []) or 'experience' in tag.text.lower()))
                st.session_state['job_experience'] = experience_element.text.strip().replace('\n', ' ') if experience_element else "Profil/Expérience non trouvé (vérifiez sélecteurs)"
                
                # For description, look for common job description containers
                description_element = soup.find(lambda tag: tag.name == 'div' and ('list-style-editor' in tag.get('class', []) or 'job-description' in tag.get('id', '') or 'description' in tag.get('id', '')))
                if not description_element: # Try to get main content if specific class not found
                    description_element = soup.find('main') or soup.find('article') or soup.find('body') # Last resort
                st.session_state['job_description'] = description_element.text.strip() if description_element else "Description du poste non trouvée (vérifiez sélecteurs)"

                st.success("Détails du poste extraits (potentiellement). Veuillez vérifier et affiner ci-dessous.")
            except requests.exceptions.RequestException as e:
                st.error(f"Erreur de réseau ou HTTP: {e}")
            except Exception as e:
                st.error(f"Erreur d'extraction URL: {e}. Vérifiez les sélecteurs ou saisissez manuellement.")
        elif submit_url and not url:
            st.warning("Veuillez entrer une URL valide.")

st.session_state['job_title'] = st.text_input("Titre de l'offre:", value=st.session_state['job_title'])
st.session_state['job_experience'] = st.text_area("Profil recherché (Expérience, Formation):", value=st.session_state['job_experience'], height=150)
st.session_state['job_description'] = st.text_area("Description du poste:", value=st.session_state['job_description'], height=300)
st.markdown("<h3 style='text-align: left;'>Téléversement des CV:</h3>", unsafe_allow_html=True)
uploaded_files = st.file_uploader("Téléchargez les CV (PDF, DOCX)", accept_multiple_files=True, type=["pdf", "docx"])
if uploaded_files:
    st.markdown(f"Total des CV téléchargés: {len(uploaded_files)}")


# --- Helper Functions for Text Extraction and Knowledge Base ---
def extract_text_from_pdf(file_obj):
    try:
        reader = PdfReader(file_obj)
        text = "".join(page.extract_text() or "" for page in reader.pages)
        return text.strip() if text else None
    except Exception as e:
        st.error(f"Erreur PDF ({getattr(file_obj, 'name', 'N/A')}): {e}")
        return None

def extract_text_from_docx(file_obj):
    try:
        doc = Document(file_obj)
        text = "\n".join(para.text for para in doc.paragraphs if para.text)
        return text.strip() if text else None
    except Exception as e:
        st.error(f"Erreur DOCX ({getattr(file_obj, 'name', 'N/A')}): {e}")
        return None

def build_knowledge_base_string():
    kb_parts = ["BASE DE CONNAISSANCES POUR LE CHAT IA:"]
    kb_parts.append("\n--- DÉTAILS DU POSTE ---")
    kb_parts.append(f"Titre du poste: {st.session_state.get('job_title', 'Non spécifié')}")
    kb_parts.append(f"Profil recherché: {st.session_state.get('job_experience', 'Non spécifié')}")
    kb_parts.append(f"Description du poste: {st.session_state.get('job_description', 'Non spécifié')}")
    if st.session_state.all_extracted_resume_texts:
        kb_parts.append("\n\n--- CONTENU DES CV ANALYSÉS ---")
        for i, resume_data in enumerate(st.session_state.all_extracted_resume_texts):
            kb_parts.append(f"\nCV #{i+1}: {resume_data['filename']}")
            kb_parts.append("--- DEBUT TEXTE CV ---")
            kb_parts.append(resume_data['text'] if resume_data['text'] else "ERREUR: Contenu du CV non extrait ou vide.")
            kb_parts.append("--- FIN TEXTE CV ---\n")
    else:
        kb_parts.append("\n\n--- CONTENU DES CV ANALYSÉS ---")
        kb_parts.append("Aucun CV n'a été traité ou le texte n'a pas pu être extrait pour la base de connaissances du chat.")
    return "\n".join(kb_parts)


# --- Main Processing Button ---
if st.button("Démarrer la Sélection et Préparer le Chat"):
    if not all([st.session_state['job_title'], st.session_state['job_experience'], st.session_state['job_description']]):
        st.error("Veuillez remplir Titre, Profil et Description du poste.")
    elif not uploaded_files:
        st.error("Veuillez télécharger au moins un CV.")
    else:
        # --- API Key Configuration Attempt ---
        api_keys_from_secrets = [
            st.secrets.get("api_keys", {}).get("key1"),
            st.secrets.get("api_keys", {}).get("key2"),
            st.secrets.get("api_keys", {}).get("key3"),
            st.secrets.get("api_keys", {}).get("key4"),
            st.secrets.get("api_keys", {}).get("key5"),
        ]
        # Filter out None values if keys are not found in secrets
        api_keys_list = [key for key in api_keys_from_secrets if key]

        if not api_keys_list:
            st.error("Aucune clé API n'a été trouvée dans les secrets Streamlit. Veuillez configurer `secrets.toml`.")
            st.stop()
            
        cv_processing_model = attempt_api_configuration(api_keys_list)

        if not cv_processing_model or not st.session_state.api_key_configured:
            st.error("La configuration de l'API a échoué. Impossible de continuer le traitement des CV.")
            # No need for st.stop() here, the 'else' block below won't execute.
        else:
            st.info("Configuration API réussie. Démarrage du traitement des CV...")
            # Reset/Initialize states for new processing run
            st.session_state.all_extracted_resume_texts = []
            st.session_state.knowledge_base_content = ""
            st.session_state.chat_messages = [] # Reset chat for new selection
            st.session_state.gemini_chat_model_instance = None
            st.session_state.gemini_chat_session_object = None
            
            MAX_REQUESTS_PER_MINUTE = 15 # Google Generative AI free tier limit (gemini-1.5-flash often 15 RPM)
            MAX_TOTAL_REQUESTS_PER_RUN = 1500 # Arbitrary limit to prevent excessive use in one go
            
            total_requests_cv_eval = 0
            request_timestamps_cv_eval = []

            results = []
            progress_bar = st.progress(0)
            status_text = st.empty()
            total_files = len(uploaded_files)

            def evaluate_resume_with_ai(resume_text, job_title, job_experience, job_description, filename, model_instance, 
                                        current_total_requests, current_timestamps_list):
                current_time = time.time()
                # Filter timestamps older than 60 seconds
                current_timestamps_list = [t for t in current_timestamps_list if current_time - t < 60]

                if len(current_timestamps_list) >= MAX_REQUESTS_PER_MINUTE:
                    # Calculate wait time based on the oldest timestamp in the current minute window
                    wait_time = 60 - (current_time - current_timestamps_list[0]) # oldest is at index 0 after filtering
                    if wait_time > 0:
                        status_text.info(f"Limite de débit (CV eval) atteinte. Attente de {wait_time:.1f}s...")
                        time.sleep(wait_time)
                    # Re-evaluate current time and filter timestamps again after waiting
                    current_time = time.time()
                    current_timestamps_list = [t for t in current_timestamps_list if current_time - t < 60]

                if current_total_requests >= MAX_TOTAL_REQUESTS_PER_RUN: # Check against per run limit
                    st.warning("Limite totale de requêtes par exécution (CV eval) atteinte.")
                    return None, current_total_requests, current_timestamps_list 

                current_total_requests += 1
                current_timestamps_list.append(time.time()) # Add timestamp for the current request
                
                prompt = f"""
                Compte tenu des exigences suivantes pour le poste:
                Titre du poste: {job_title}
                Exigences d'expérience et profil recherché: {job_experience}
                Description du poste: {job_description}

                Et de ce CV (nom de fichier: {filename}):
                ---CV START---
                {resume_text}
                ---CV END---

                Analysez en profondeur le CV. Veuillez extraire les informations suivantes et les présenter sous forme de texte brut au format JSON. NE PAS utiliser de blocs de code JSON (```json ... ```).
                Les champs doivent être remplis même si l'information n'est pas explicitement trouvée (utiliser "N/A" ou "Non spécifié").
                - Nom du candidat: (Nom complet)
                - Adresse e-mail: (email)
                - Numéro de téléphone: (numéro)
                - Ville: (Ville de résidence actuelle si trouvée, sinon N/A)
                - Pays: (Pays de résidence actuelle si trouvé, sinon N/A)
                - Nombre total d'années d'expérience pertinente: (Calculer le nombre d'années d'expérience professionnelle pertinente pour le poste. Exclure les stages et les expériences non pertinentes. Fournir un nombre entier ou "N/A".)
                - Pourcentage d'admissibilité au poste: (Estimer un pourcentage de 0 à 100% basé sur l'adéquation globale avec les exigences. Doit être un nombre suivi de '%'. Exemple: "85%")
                - Commentaires: (Synthèse de 3-5 phrases justifiant le pourcentage, mentionnant les forces et faiblesses par rapport aux exigences. Soyez spécifique.)
                - Sexe: (Homme, Femme, Non spécifié)
                - Formation: (Plus haut niveau de diplôme pertinent. Ex: Bac+5, Master, Licence, Bac, etc. ou N/A)
                - Date de naissance: (AAAA-MM-JJ si trouvée, sinon N/A)
                - Questions d'entretien: (Générez une liste de 5 à 7 questions CIBLÉES et PERTINENTES à poser au candidat pour approfondir les aspects clés de son profil en lien avec le poste. Chaque question doit être sur une nouvelle ligne, sans numérotation initiale.)

                Exemple de format de réponse (texte brut, pas de JSON réel):
                {{
                  "Nom du candidat": "John Doe",
                  "Adresse e-mail": "john.doe@example.com",
                  "Numéro de téléphone": "+15551234567",
                  "Ville": "Paris",
                  "Pays": "France",
                  "Nombre total d'années d'expérience pertinente": "5",
                  "Pourcentage d'admissibilité": "85%",
                  "Commentaires": "John possède une solide expérience en gestion de projet et correspond bien aux exigences. Cependant, il manque d'expérience spécifique dans la technologie X, ce qui est un point à clarifier. Ses compétences en communication semblent être un atout.",
                  "Sexe": "Homme",
                  "Formation": "Bac+5",
                  "Date de naissance": "1990-01-01",
                  "Questions d'entretien": [
                    "Pouvez-vous décrire un projet où vous avez utilisé la méthodologie Agile et quels ont été les défis rencontrés?",
                    "Comment gérez-vous les conflits au sein d'une équipe projet?",
                    "Quelle est votre expérience avec l'outil Y mentionné dans notre description de poste?",
                    "Décrivez une situation où vous avez dû prendre une décision difficile avec des informations limitées.",
                    "Comment vous tenez-vous informé des dernières tendances dans [domaine du poste]?"
                  ]
                }}
                Veuillez répondre en français.
                """
                try:
                    response = model_instance.generate_content(prompt)
                    return response.text.strip(), current_total_requests, current_timestamps_list
                except Exception as e:
                    st.error(f"Erreur API Gemini pour {filename}: {e}")
                    if "quota" in str(e).lower() or "rate limit" in str(e).lower(): # More specific check
                        st.warning(f"Quota API ou limite de débit potentiellement atteint pour {filename}. L'erreur était: {e}")
                        # Consider adding a longer sleep here or stopping if it's a hard quota error
                    return None, current_total_requests, current_timestamps_list

            for i, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"Traitement CV {i+1}/{total_files}: {uploaded_file.name}")
                filename = uploaded_file.name
                file_extension = os.path.splitext(filename)[-1].lower().replace(".", "") # More robust extension extraction
                text_content = None

                try:
                    if file_extension == "pdf":
                        text_content = extract_text_from_pdf(io.BytesIO(uploaded_file.getvalue()))
                    elif file_extension == "docx":
                        text_content = extract_text_from_docx(io.BytesIO(uploaded_file.getvalue()))
                    else:
                        st.warning(f"Type de fichier non supporté: {filename}")
                        results.append({"Nom du fichier": filename, "Nom du candidat": "N/A", "Admissibilité (%)": 0, "Commentaires": "Type non supporté", "Questions d'entretien": [], "Job Title": st.session_state['job_title'], "Gender": "N/A", "Formation": "N/A", "Ville": "N/A", "Pays": "N/A", "Expérience (Années)": "N/A", "Date de naissance": "N/A", "Téléphone": "N/A", "E-mail": "N/A"})
                        st.session_state.all_extracted_resume_texts.append({'filename': filename, 'text': "Type de fichier non supporté."})
                        progress_bar.progress((i + 1) / total_files)
                        continue
                except Exception as e:
                    st.error(f"Erreur lecture fichier {filename}: {e}")
                    results.append({"Nom du fichier": filename, "Nom du candidat": "N/A", "Admissibilité (%)": 0, "Commentaires": f"Erreur lecture: {e}", "Questions d'entretien": [], "Job Title": st.session_state['job_title'], "Gender": "N/A", "Formation": "N/A", "Ville": "N/A", "Pays": "N/A", "Expérience (Années)": "N/A", "Date de naissance": "N/A", "Téléphone": "N/A", "E-mail": "N/A"})
                    st.session_state.all_extracted_resume_texts.append({'filename': filename, 'text': f"Erreur de lecture du fichier: {e}"})
                    progress_bar.progress((i + 1) / total_files)
                    continue
                
                if total_requests_cv_eval >= MAX_TOTAL_REQUESTS_PER_RUN: 
                    st.warning("Limite maximale du nombre total de requêtes CV eval pour cette exécution atteinte. Arrêt du traitement des CV.")
                    break

                if text_content:
                    st.session_state.all_extracted_resume_texts.append({'filename': filename, 'text': text_content})
                    
                    ai_response, total_requests_cv_eval, request_timestamps_cv_eval = evaluate_resume_with_ai(
                        text_content, 
                        st.session_state['job_title'], 
                        st.session_state['job_experience'], 
                        st.session_state['job_description'], 
                        filename, 
                        cv_processing_model, # This is the model from attempt_api_configuration
                        total_requests_cv_eval, 
                        request_timestamps_cv_eval 
                    )
                    
                    if ai_response:
                        try:
                            # Using re.search with robust error handling for each field
                            def get_match_group(pattern, text, group=1, default="N/A"):
                                match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                                return match.group(group).strip() if match and match.group(group) else default

                            candidate_name_val = get_match_group(r'"Nom du candidat":\s*"([^"]*)"', ai_response).upper()
                            email_val = get_match_group(r'"Adresse e-mail":\s*"([^"]*)"', ai_response)
                            phone_val = get_match_group(r'"Numéro de téléphone":\s*"([^"]*)"', ai_response)
                            city_val = get_match_group(r'"Ville":\s*"([^"]*)"', ai_response)
                            country_val = get_match_group(r'"Pays":\s*"([^"]*)"', ai_response)
                            experience_val = get_match_group(r'"Nombre total d\'années d\'expérience pertinente":\s*"([^"]*)"', ai_response)
                            
                            admissibility_str_match = re.search(r'"Pourcentage d\'admissibilité":\s*"([\d]+)\s*%"', ai_response, re.IGNORECASE)
                            admissibility_val = int(admissibility_str_match.group(1)) if admissibility_str_match else 0
                            
                            comments_val = get_match_group(r'"Commentaires":\s*"((?:[^"\\]|\\.)*)"', ai_response).replace('\\"', '"')
                            gender_val = get_match_group(r'"Sexe":\s*"([^"]*)"', ai_response, default="Non spécifié")
                            formation_val = get_match_group(r'"Formation":\s*"([^"]*)"', ai_response)
                            date_naissance_val = get_match_group(r'"Date de naissance":\s*"([^"]*)"', ai_response)

                            questions_list = []
                            questions_block_match = re.search(r'"Questions d\'entretien":\s*\[\s*([\s\S]*?)\s*\]', ai_response, re.IGNORECASE)
                            if questions_block_match:
                                questions_str = questions_block_match.group(1)
                                # Split by newline, then process each line
                                raw_questions = questions_str.split('\n')
                                for q_line in raw_questions:
                                    # Remove leading/trailing whitespace, quotes, and commas
                                    q_line_stripped = q_line.strip().strip(',').strip('"').strip()
                                    if q_line_stripped:
                                        questions_list.append(q_line_stripped)
                            
                            new_filename = f"{admissibility_val}% - {candidate_name_val if candidate_name_val != 'N/A' else os.path.splitext(filename)[0]}.{file_extension}"

                            results.append({
                                "Nom du fichier": new_filename, "Nom du candidat": candidate_name_val,
                                "Job Title": st.session_state['job_title'], "Admissibilité (%)": admissibility_val,
                                "Commentaires": comments_val, "Questions d'entretien": questions_list,
                                "Gender": gender_val, "Formation": formation_val, "Ville": city_val,
                                "Pays": country_val, "Expérience (Années)": experience_val,
                                "Date de naissance": date_naissance_val, "Téléphone": phone_val, "E-mail": email_val
                            })
                        except Exception as e:
                            st.error(f"Erreur parsing IA pour {filename}: {e}. Réponse: {ai_response[:300]}...")
                            results.append({"Nom du fichier": filename, "Nom du candidat": "Erreur Analyse", "Admissibilité (%)": 0, "Commentaires": f"Erreur analyse IA: {e}", "Questions d'entretien": [], "Job Title": st.session_state['job_title'], "Gender": "N/A", "Formation": "N/A", "Ville": "N/A", "Pays": "N/A", "Expérience (Années)": "N/A", "Date de naissance": "N/A", "Téléphone": "N/A", "E-mail": "N/A"})
                    else:
                        st.warning(f"Pas de réponse IA pour {filename} (peut-être dû à une erreur API précédente).")
                        results.append({"Nom du fichier": filename, "Nom du candidat": "Pas réponse IA", "Admissibilité (%)": 0, "Commentaires": "Pas de réponse IA", "Questions d'entretien": [], "Job Title": st.session_state['job_title'], "Gender": "N/A", "Formation": "N/A", "Ville": "N/A", "Pays": "N/A", "Expérience (Années)": "N/A", "Date de naissance": "N/A", "Téléphone": "N/A", "E-mail": "N/A"})
                else:
                    st.warning(f"Texte non extrait de {filename} ou fichier vide.")
                    results.append({"Nom du fichier": filename, "Nom du candidat": "Erreur Texte", "Admissibilité (%)": 0, "Commentaires": "Texte non extrait", "Questions d'entretien": [], "Job Title": st.session_state['job_title'], "Gender": "N/A", "Formation": "N/A", "Ville": "N/A", "Pays": "N/A", "Expérience (Années)": "N/A", "Date de naissance": "N/A", "Téléphone": "N/A", "E-mail": "N/A"})
                    st.session_state.all_extracted_resume_texts.append({'filename': filename, 'text': "ERREUR: Échec de l'extraction du texte."})
                progress_bar.progress((i + 1) / total_files)

            status_text.success("Sélection des CV terminée!")

            if results:
                df = pd.DataFrame(results)
                expected_columns = ["Nom du fichier", "Nom du candidat", "Job Title", "Admissibilité (%)", "Commentaires", "Questions d'entretien", "Gender",
                                    "Formation", "Ville", "Pays", "Expérience (Années)", "Date de naissance", "Téléphone", "E-mail"]
                # Ensure all expected columns exist, fill with defaults if not
                for col in expected_columns:
                    if col not in df.columns:
                        if col == "Questions d'entretien":
                            df[col] = [[] for _ in range(len(df))]
                        elif col == "Admissibilité (%)":
                            df[col] = 0
                        else:
                            df[col] = "N/A"
                
                df = df[expected_columns] # Reorder and select columns
                df = df.sort_values(by="Admissibilité (%)", ascending=False).reset_index(drop=True)
                
                # Format questions for display in dataframe
                df_display = df.copy()
                df_display["Questions d'entretien"] = df_display["Questions d'entretien"].apply(lambda x: "\n".join(x) if isinstance(x, list) else x)
                st.dataframe(df_display) # Removed style for simplicity, can be added back if needed

                # --- Prepare knowledge base and chat model ---
                st.session_state.knowledge_base_content = build_knowledge_base_string()
                
                if st.session_state.knowledge_base_content and st.session_state.api_key_configured:
                    st.info("Préparation du module de chat avec la base de connaissances...")
                    system_instruction_text = f"""
                    Vous êtes un assistant IA expert en ressources humaines. Votre unique source de vérité est la base de connaissances suivante.
                    Ne répondez JAMAIS en utilisant des informations extérieures à cette base.

                    --- DÉBUT BASE DE CONNAISSANCES ---
                    {st.session_state.knowledge_base_content}
                    --- FIN BASE DE CONNAISSANCES ---

                    Lorsque vous répondez aux questions de l'utilisateur, basez-vous EXCLUSIVEMENT sur les CV, la description du poste, le titre du poste et l'expérience requise fournis dans cette base de connaissances.
                    Si une question ne peut pas être répondue à partir de ces informations, indiquez-le clairement (par exemple: "L'information demandée n'est pas disponible dans les documents fournis.").
                    Soyez concis et précis. Répondez en français.
                    """
                    try:
                        # The API key is already configured globally by attempt_api_configuration
                        st.session_state.gemini_chat_model_instance = genai.GenerativeModel(
                            model_name='gemini-1.5-flash',
                            system_instruction=system_instruction_text
                        )
                        st.session_state.gemini_chat_session_object = st.session_state.gemini_chat_model_instance.start_chat(history=[]) # Start with empty history for new selection
                        st.session_state.chat_messages.append( # Initial message for this new chat session
                            {"role": "assistant", "content": "Bonjour ! La base de connaissances a été chargée pour cette sélection. Comment puis-je vous aider à analyser ces CVs et cette offre d'emploi ?"}
                        )
                        st.success("Module de chat prêt. Vous pouvez maintenant poser des questions ci-dessous.")
                    except Exception as e:
                        st.error(f"Erreur lors de l'initialisation du modèle de chat Gemini : {e}")
                        st.session_state.gemini_chat_model_instance = None
                        st.session_state.gemini_chat_session_object = None
                elif not st.session_state.api_key_configured:
                     st.warning("L'API n'est pas configurée. Le module de chat ne peut pas être initialisé.")
                else: # Knowledge base content might be empty if no CVs processed successfully
                    st.warning("La base de connaissances pour le chat est vide ou n'a pas pu être construite.")


                # --- DOWNLOAD BUTTONS ---
                def to_excel(df_to_export):
                    output = io.BytesIO()
                    df_excel = df_to_export.copy()
                    if "Questions d'entretien" in df_excel.columns:
                         df_excel["Questions d'entretien"] = df_excel["Questions d'entretien"].apply(lambda x: "\n".join(x) if isinstance(x, list) else x)
                    writer = pd.ExcelWriter(output, engine='xlsxwriter')
                    df_excel.to_excel(writer, sheet_name='Resultats CV', index=False)
                    worksheet = writer.sheets['Resultats CV']
                    for idx, col_name_excel in enumerate(df_excel.columns): # Iterate over actual columns in df_excel
                        series = df_excel[col_name_excel]
                        # Calculate max_len based on content and header
                        max_len_content = series.astype(str).map(len).max()
                        if pd.isna(max_len_content): max_len_content = 0 # Handle case where all values are NaN
                        
                        max_len = max(max_len_content, len(str(series.name))) + 2 # Add a little padding
                        if max_len > 60: max_len = 60 # Cap max width
                        worksheet.set_column(idx, idx, max_len)
                    writer.close() # Use close() for xlsxwriter
                    return output.getvalue()

                def set_portrait(document):
                    section = document.sections[0]
                    if section.orientation != WD_ORIENT.PORTRAIT:
                        section.orientation = WD_ORIENT.PORTRAIT
                        new_width, new_height = section.page_height, section.page_width
                        section.page_width = new_width; section.page_height = new_height
                    # Standard A4 Portrait
                    section.page_width = Inches(8.27); section.page_height = Inches(11.69)
                    # Set margins (optional, but good for layout)
                    section.left_margin = Inches(0.75)
                    section.right_margin = Inches(0.75)
                    section.top_margin = Inches(0.75)
                    section.bottom_margin = Inches(0.75)


                def create_styled_docx(df_to_export, job_title_param):
                    document = Document()
                    set_portrait(document)
                    
                    title_paragraph = document.add_paragraph(); title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title_paragraph.add_run(f"Sélection CV pour : {job_title_param}"); title_run.bold = True; title_run.font.size = Pt(16)
                    document.add_paragraph() 
                    
                    fields_to_display = [
                        ("Job Title", "Emploi"), ("Admissibilité (%)", "Admissibilité"), ("Commentaires", "Commentaires"),
                        ("Gender", "Sexe"), ("Formation", "Formation"), ("Ville", "Ville"), ("Pays", "Pays"),
                        ("Expérience (Années)", "Expérience (Années)"), ("Date de naissance", "Date de naissance"),
                        ("Téléphone", "Téléphone"), ("E-mail", "E-mail")
                    ] # Questions handled separately

                    for index, row_data in df_to_export.iterrows():
                        if index > 0: document.add_page_break()
                        
                        candidate_name = str(row_data.get("Nom du candidat", "N/A")).upper()
                        p_name = document.add_paragraph(); p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_name = p_name.add_run(candidate_name); run_name.bold = True; run_name.font.size = Pt(14)
                        document.add_paragraph() # Spacer

                        # Table for candidate details
                        if any(row_data.get(col_name_docx) not in [None, "N/A", ""] for col_name_docx, _ in fields_to_display):
                            table = document.add_table(rows=1, cols=2) # Start with 1 row for headers (optional, or just data)
                            table.style = 'TableGrid' # Add a visible grid
                            table.autofit = False
                            table.columns[0].width = Inches(2.2) 
                            table.columns[1].width = Inches(5.0) 
                            
                            # Remove header row if not needed, or populate it. For this layout, direct data entry is fine.
                            # For now, let's assume no explicit header row in the table, just label-value pairs.
                            # So, we'll add rows dynamically.
                            table.rows[0]._element.getparent().remove(table.rows[0]._element) # Remove initial blank row

                            for col_name_docx, display_label in fields_to_display:
                                value_to_display = row_data.get(col_name_docx, "N/A")
                                if col_name_docx == "Job Title": value_to_display = job_title_param # Use actual job title
                                if col_name_docx == "Admissibilité (%)": value_to_display = f"{value_to_display}%"

                                if value_to_display not in [None, "N/A", ""] or col_name_docx == "Job Title": # Show job title always
                                    cells = table.add_row().cells
                                    p_label = cells[0].paragraphs[0]; run_label = p_label.add_run(f"{display_label} :"); run_label.bold = True
                                    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    for run in p_label.runs: run.font.size = Pt(10)
                                    
                                    cells[1].text = str(value_to_display)
                                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    for run in cells[1].paragraphs[0].runs: run.font.size = Pt(10)
                        
                        document.add_paragraph() # Spacer
                        questions_label_p = document.add_paragraph(); run_q_label = questions_label_p.add_run("Questions d'entretien suggérées :"); run_q_label.bold = True; run_q_label.font.size = Pt(11)
                        questions_data = row_data.get("Questions d'entretien", [])
                        if isinstance(questions_data, list) and questions_data:
                            for question_text in questions_data:
                                cleaned_question = re.sub(r"^\s*[\d\-\*\.]+\s*", "", str(question_text).strip()) # Remove list markers
                                if cleaned_question: 
                                    p_q = document.add_paragraph(cleaned_question, style='ListBullet')
                                    for run in p_q.runs: run.font.size = Pt(10)
                        elif isinstance(questions_data, str) and questions_data.strip() not in ["N/A", ""]: 
                            p_q = document.add_paragraph(questions_data.strip(), style='ListBullet')
                            for run in p_q.runs: run.font.size = Pt(10)
                        else: 
                            p_q = document.add_paragraph("N/A", style='ListBullet') # Or just "Aucune question spécifique générée."
                            for run in p_q.runs: run.font.size = Pt(10)

                    # Summary Table at the end
                    document.add_page_break()
                    summary_title_p = document.add_paragraph(); summary_title_run = summary_title_p.add_run("Tableau Récapitulatif des Candidats"); summary_title_run.bold = True; summary_title_run.font.size = Pt(14); summary_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph()
                    
                    df_sorted_summary = df_to_export.sort_values(by="Admissibilité (%)", ascending=False).reset_index(drop=True)
                    summary_table = document.add_table(rows=1, cols=3); summary_table.style = 'TableGrid'; summary_table.autofit = False
                    summary_table.columns[0].width = Inches(2.5); summary_table.columns[1].width = Inches(1.0); summary_table.columns[2].width = Inches(3.5) # Adjusted widths
                    
                    hdr_cells = summary_table.rows[0].cells
                    hdr_cells[0].text = "Nom du Candidat"; hdr_cells[1].text = "Admissibilité"; hdr_cells[2].text = "Commentaires Clés"
                    for cell in hdr_cells: 
                        if cell.paragraphs:
                            cell.paragraphs[0].runs[0].bold = True
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in cell.paragraphs[0].runs: run.font.size = Pt(10)
                    
                    for _, row in df_sorted_summary.iterrows():
                        row_cells = summary_table.add_row().cells
                        row_cells[0].text = str(row.get("Nom du candidat", "N/A"))
                        row_cells[1].text = str(row.get("Admissibilité (%)", "0")) + "%"
                        comments_summary = str(row.get("Commentaires", "N/A"))
                        row_cells[2].text = comments_summary[:147] + "..." if len(comments_summary) > 150 else comments_summary
                        for cell_idx, cell in enumerate(row_cells):
                            if cell.paragraphs:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs: run.font.size = Pt(9)
                                if cell_idx == 1: # Center admissibility percentage
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                    docx_stream = io.BytesIO(); document.save(docx_stream); docx_stream.seek(0); return docx_stream

                excel_file = to_excel(df) # Use the original df for Excel
                st.download_button(
                    label="Télécharger les résultats au format Excel",
                    data=excel_file,
                    file_name=f'resultats_selection_cv_{st.session_state["job_title"].replace(" ", "_")[:30]}.xlsx',
                    mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                )

                docx_file = create_styled_docx(df, st.session_state['job_title']) # Use original df for Word
                st.download_button(
                    label="Télécharger les résultats au format Word (DOCX)",
                    data=docx_file,
                    file_name=f'rapport_selection_cv_{st.session_state["job_title"].replace(" ", "_")[:30]}.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            else: # This 'else' corresponds to 'if results:'
                st.warning("Aucun résultat à afficher ou à exporter (aucun CV traité avec succès ou aucune donnée extraite).")


# --- Chat Interface Section ---
if 'gemini_chat_session_object' in st.session_state and st.session_state.gemini_chat_session_object:
    st.markdown("---")
    st.subheader("💬 Discuter avec l'IA à propos des CVs et du Poste")
    num_cvs_for_chat = len(st.session_state.get('all_extracted_resume_texts', []))
    st.markdown(f"<small>L'IA a connaissance du titre du poste, de la description, du profil recherché, et du contenu des {num_cvs_for_chat} CVs traités lors de la dernière sélection.</small>", unsafe_allow_html=True)
    
    # Display chat messages
    for message in st.session_state.chat_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            
    if user_prompt := st.chat_input("Posez votre question sur les CVs ou le poste..."):
        st.session_state.chat_messages.append({"role": "user", "content": user_prompt})
        with st.chat_message("user"):
            st.markdown(user_prompt)
        
        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            message_placeholder.markdown("🧠 L'IA réfléchit...")
            try:
                # Add a small delay if desired, e.g., time.sleep(0.5)
                response = st.session_state.gemini_chat_session_object.send_message(user_prompt)
                
                full_response_text = ""
                # Handle potential streaming or different response structures
                if hasattr(response, 'text') and response.text:
                    full_response_text = response.text
                elif hasattr(response, 'parts'):
                    for part in response.parts: # Iterate if it's a list of parts
                        if hasattr(part, 'text'):
                            full_response_text += part.text
                elif isinstance(response, str): # If response is directly a string
                     full_response_text = response
                else: # Fallback if structure is unknown
                    full_response_text = "Réponse de l'IA reçue, mais format inattendu."
                    st.warning(f"Format de réponse inattendu: {type(response)}")

                message_placeholder.markdown(full_response_text)
                st.session_state.chat_messages.append({"role": "assistant", "content": full_response_text})
            except Exception as e:
                error_msg = f"Désolé, une erreur est survenue avec l'IA de chat : {e}"
                message_placeholder.error(error_msg)
                st.session_state.chat_messages.append({"role": "assistant", "content": error_msg})
elif 'knowledge_base_content' in st.session_state and st.session_state.knowledge_base_content and st.session_state.get('api_key_configured', False) and not ('gemini_chat_session_object' in st.session_state and st.session_state.gemini_chat_session_object):
    st.warning("Le module de chat n'a pas pu être initialisé correctement après le traitement des CV. Vérifiez les messages d'erreur ci-dessus. L'API semble configurée, mais l'objet de chat est manquant.")
elif not st.session_state.get('api_key_configured', False) and st.button("Vérifier la Configuration API pour le Chat (si déjà traité)"):
    # This button is a bit redundant if API config happens with the main button,
    # but could be a manual trigger if something went wrong post-processing.
    # It might be better to just rely on the main button's API setup.
    st.info("Tentative de configuration de l'API pour le chat (si les CVs ont déjà été traités et que vous voulez activer le chat maintenant)...")
    api_keys_from_secrets = [
            st.secrets.get("api_keys", {}).get("key1"),
            st.secrets.get("api_keys", {}).get("key2"), # ... etc
        ]
    api_keys_list = [key for key in api_keys_from_secrets if key]
    if api_keys_list:
        test_model = attempt_api_configuration(api_keys_list)
        if test_model and st.session_state.api_key_configured and st.session_state.knowledge_base_content:
            # Re-attempt chat initialization (copy logic from above)
            st.info("Préparation du module de chat...")
            # ... (copy system instruction and model init)
            # This part might be complex to re-trigger cleanly without reprocessing.
            # It's generally better to ensure chat is set up immediately after CV processing.
            st.success("API configurée. Si le chat ne s'affiche pas, réessayez de traiter les CVs.")
        else:
            st.error("Échec de la configuration de l'API pour le chat.")
    else:
        st.error("Aucune clé API trouvée dans les secrets pour la configuration du chat.")
